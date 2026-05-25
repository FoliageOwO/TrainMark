#!/usr/bin/env python3
"""PaddleOCR-backed OCR provider for TrainMark AI.

The script keeps the TrainMark OCR JSON contract stable. When PaddleOCR or the
input artifact is unavailable in a local MVP environment, it falls back to the
same deterministic extraction shape used by the local provider unless
--require-real is set.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[2]
os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
os.environ.setdefault("PADDLE_PDX_ENABLE_MKLDNN_BYDEFAULT", "False")
os.environ.setdefault("FLAGS_use_mkldnn", "false")


@dataclass(frozen=True)
class OcrBlock:
    type: str
    title: str
    page: int
    confidence: int


def infer_blocks(object_key: str) -> list[OcrBlock]:
    normalized = object_key.lower()
    if "database" in normalized or "数据库" in normalized:
        return [
            OcrBlock("heading", "数据库概念结构设计", 1, 95),
            OcrBlock("table", "ER 实体关系表", 3, 92),
            OcrBlock("table", "数据字典", 5, 90),
            OcrBlock("paragraph", "规范化分析", 7, 88),
        ]
    if normalized.endswith((".png", ".jpg", ".jpeg")):
        return [
            OcrBlock("image", "系统截图", 1, 89),
            OcrBlock("paragraph", "截图文字说明", 1, 84),
        ]
    return [
        OcrBlock("heading", "需求分析", 2, 96),
        OcrBlock("table", "数据库表结构", 7, 91),
        OcrBlock("image", "系统运行截图", 12, 88),
        OcrBlock("heading", "实训总结", 17, 90),
    ]


def build_plain_text_preview(blocks: list[OcrBlock], source: str) -> str:
    titles = "、".join(block.title for block in blocks) or "文档内容"
    return f"识别到 {titles} 等结构化内容。来源：{source}。"


def configured_upload_root() -> Path:
    root = Path(os.environ.get("UPLOAD_OBJECT_ROOT", ".data/uploads"))
    if root.is_absolute():
        return root
    return ROOT / root


def resolve_object_path(normalized_object_key: str, object_key: str) -> Path:
    values = [value for value in (normalized_object_key, object_key) if value]
    upload_root = configured_upload_root()
    for value in values:
        candidate = Path(value)
        candidates = [
            candidate,
            ROOT / candidate,
            upload_root / candidate,
        ]
        for path in candidates:
            if path.is_file():
                return path
    fallback = normalized_object_key or object_key
    return upload_root / fallback if fallback else upload_root


def resolve_input_path(args: argparse.Namespace) -> Path:
    return resolve_object_path(args.normalized_object_key, args.object_key)


def result_payload(job_id: int, submission_id: int, blocks: list[OcrBlock], source: str) -> dict[str, Any]:
    return {
        "jobId": job_id,
        "submissionId": submission_id,
        "plainText": build_plain_text_preview(blocks, source),
        "blocks": [asdict(block) for block in blocks],
    }


def normalize_result_item(item: Any) -> dict[str, Any]:
    if isinstance(item, dict):
        return item.get("res", item)
    res = getattr(item, "res", None)
    if isinstance(res, dict):
        return res
    json_data = getattr(item, "json", None)
    if isinstance(json_data, dict):
        return json_data.get("res", json_data)
    return {}


def blocks_from_paddle_result(result: Any) -> list[OcrBlock]:
    blocks: list[OcrBlock] = []
    for page_number, item in enumerate(result, start=1):
        data = normalize_result_item(item)
        page = data.get("page_index")
        if isinstance(page, int):
            page_number = page + 1

        texts = data.get("rec_texts")
        scores = data.get("rec_scores")
        if isinstance(texts, list):
            for index, text in enumerate(texts):
                title = str(text).strip()
                if not title:
                    continue
                confidence = score_to_percent(scores[index] if isinstance(scores, list) and index < len(scores) else None)
                blocks.append(OcrBlock("paragraph", title, page_number, confidence))
            continue

        text = data.get("rec_text")
        if isinstance(text, str) and text.strip():
            blocks.append(OcrBlock("paragraph", text.strip(), page_number, score_to_percent(data.get("rec_score"))))
    return blocks


def score_to_percent(score: Any) -> int:
    try:
        value = float(score)
    except (TypeError, ValueError):
        return 90
    if value <= 1:
        value *= 100
    return max(0, min(100, round(value)))


def recognize_with_paddle(input_path: Path, args: argparse.Namespace) -> list[OcrBlock]:
    from paddleocr import PaddleOCR  # type: ignore

    options = {
        "use_doc_orientation_classify": False,
        "use_doc_unwarping": False,
        "use_textline_orientation": False,
        "lang": args.language,
        "device": os.environ.get("OCR_DEVICE", "cpu"),
        "enable_mkldnn": env_bool("OCR_ENABLE_MKLDNN", False),
        "cpu_threads": int(os.environ.get("OCR_CPU_THREADS", "4")),
        "engine": args.engine,
    }
    ocr = build_paddleocr(PaddleOCR, options)

    result = ocr.predict(str(input_path))
    return blocks_from_paddle_result(result)


def is_text_document(input_path: Path) -> bool:
    return input_path.suffix.lower() in {".docx", ".pdf", ".txt", ".md"}


def extract_text_document_blocks(input_path: Path) -> list[OcrBlock]:
    suffix = input_path.suffix.lower()
    if suffix == ".docx":
        return docx_blocks(input_path)
    if suffix == ".pdf":
        blocks = pdf_text_blocks(input_path)
        if blocks:
            return blocks
    if suffix in {".txt", ".md"}:
        return text_file_blocks(input_path)
    return []


def docx_blocks(input_path: Path) -> list[OcrBlock]:
    from docx import Document  # type: ignore

    document = Document(str(input_path))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return lines_to_blocks(lines)


def pdf_text_blocks(input_path: Path) -> list[OcrBlock]:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        return []

    reader = PdfReader(str(input_path))
    blocks: list[OcrBlock] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        for line in split_meaningful_lines(text):
            blocks.append(OcrBlock("paragraph", line, page_index, 95))
    return blocks


def text_file_blocks(input_path: Path) -> list[OcrBlock]:
    return lines_to_blocks(input_path.read_text(encoding="utf-8", errors="ignore").splitlines())


def lines_to_blocks(lines: list[str]) -> list[OcrBlock]:
    blocks: list[OcrBlock] = []
    for line in lines:
        for text in split_meaningful_lines(line):
            block_type = "heading" if len(text) <= 32 and not any(mark in text for mark in "。，；,.") else "paragraph"
            blocks.append(OcrBlock(block_type, text, 1, 95))
    return blocks


def split_meaningful_lines(text: str) -> list[str]:
    return [line.strip() for line in text.splitlines() if line.strip()]


def build_paddleocr(factory: Any, options: dict[str, Any]) -> Any:
    remaining = dict(options)
    for _ in range(len(options) + 1):
        try:
            return factory(**remaining)
        except (TypeError, ValueError) as error:
            unsupported = unsupported_argument(str(error), remaining)
            if unsupported is None:
                raise
            remaining.pop(unsupported, None)
    return factory(**remaining)


def unsupported_argument(message: str, options: dict[str, Any]) -> str | None:
    match = re.search(r"Unknown argument: ([A-Za-z0-9_]+)", message)
    if match and match.group(1) in options:
        return match.group(1)
    for name in options:
        if name in message:
            return name
    return None


def env_bool(name: str, default: bool) -> bool:
    value = os.environ.get(name)
    if value is None:
        return default
    return value.lower() in {"1", "true", "yes", "on"}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run PaddleOCR and emit TrainMark OCR JSON.")
    parser.add_argument("--job-id", type=int, required=True)
    parser.add_argument("--submission-id", type=int, required=True)
    parser.add_argument("--object-key", required=True)
    parser.add_argument("--normalized-object-key", default="")
    parser.add_argument("--language", default="ch")
    parser.add_argument("--engine", default="paddle", choices=["paddle", "transformers"])
    parser.add_argument("--require-real", action="store_true", help="PaddleOCR 不可用时直接失败，不使用离线兜底")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    input_path = resolve_input_path(args)
    blocks: list[OcrBlock]
    source = "PaddleOCR"

    if input_path.exists():
        try:
            if is_text_document(input_path):
                blocks = extract_text_document_blocks(input_path)
                source = "文档文本提取"
            else:
                blocks = recognize_with_paddle(input_path, args)
        except Exception as error:  # noqa: BLE001 - provider boundary logs and falls back.
            if args.require_real:
                raise RuntimeError(f"PaddleOCR 必须可用，但当前调用失败：{error}") from error
            print(f"[paddleocr-provider] PaddleOCR 不可用，使用离线兜底：{error}", file=sys.stderr)
            source = "PaddleOCR 离线兜底"
            blocks = infer_blocks(args.object_key)
    else:
        if args.require_real:
            raise FileNotFoundError(f"PaddleOCR 输入文件不存在：{input_path}")
        print(f"[paddleocr-provider] 输入文件不存在，使用离线兜底：{input_path}", file=sys.stderr)
        source = "PaddleOCR 离线兜底"
        blocks = infer_blocks(args.object_key)

    if not blocks:
        if args.require_real:
            raise RuntimeError("PaddleOCR 没有返回可用文本块")
        source = "PaddleOCR 离线兜底"
        blocks = infer_blocks(args.object_key)

    print(json.dumps(result_payload(args.job_id, args.submission_id, blocks, source), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
